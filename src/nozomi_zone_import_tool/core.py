"""
Core, GUI-independent logic for the Nozomi Guardian Zone Import Tool:
field schema, file loading (CSV/Excel/Word), column auto-mapping, Guardian
config generation/parsing, and deriving suggested zones from an asset
export. Kept free of any GUI imports so it can be unit tested and reused
from the CLI without needing Tkinter installed.
"""

import csv
import difflib
import os
import re

import pandas as pd

try:
    import docx  # python-docx
except ImportError:
    docx = None


# --------------------------------------------------------------------------
# Guardian zone field schema
# --------------------------------------------------------------------------
# key                                   label                                    cli command                              always_written_if_full  never_written_if_empty
FIELDS = [
    ("zone_name",                           "Zone Name",                              None,                                       True,  False),
    ("networks",                            "Networks / CIDR (comma separated)",      None,                                       True,  False),
    ("level",                               "Level",                                  "setlevel",                                 True,  False),
    ("security_profile",                    "Security Profile",                       "setsecurity_profile",                      True,  False),
    ("force_assigned_vlan_id",              "Force Assigned VLAN ID",                 "setforce_assigned_vlan_id",                True,  False),
    ("mac_matching_fallback",               "MAC Matching Fallback",                  "setmac_matching_fallback",                 True,  False),
    ("extended_network_statistic_enabled",  "Extended Network Statistic Enabled",     "setextended_network_statistic_enabled",    True,  False),
    ("use_label_as_node_device_id",         "Use Label As Node Device ID",            "setuse_label_as_node_device_id",           True,  False),
    ("matching_vlan_id",                    "Matching VLAN ID",                       "setmatching_vlan_id",                      True,  False),
    ("assigned_vlan_id",                    "Assigned VLAN ID",                       "setassigned_vlan_id",                      True,  False),
    ("adaptive_learning",                   "Adaptive Learning",                      "setadaptive_learning",                     True,  False),
    ("learning",                            "Learning",                               "setlearning",                              True,  False),
    ("is_public",                           "Is Public",                              "setis_public",                             True,  False),
    ("merging_zone",                        "Merging Zone",                           "setmerging_zone",                          False, True),
    ("merging",                             "Merging",                                "setmerging",                               True,  False),
    ("isolated",                            "Isolated",                               "setisolated",                              True,  False),
]
FIELD_KEYS = [f[0] for f in FIELDS]
FIELD_LABELS = {f[0]: f[1] for f in FIELDS}
FIELD_CMDS = {f[0]: f[2] for f in FIELDS if f[2]}
FIELD_ALWAYS_FULL = {f[0]: f[3] for f in FIELDS}
FIELD_NEVER_IF_EMPTY = {f[0]: f[4] for f in FIELDS}
EDITABLE_FIELD_KEYS = FIELD_KEYS + ["output_mode"]  # output_mode is a per-row override

# Synonyms used for automatic column -> field mapping
SYNONYMS = {
    "zone_name": ["zone", "zone name", "name", "zonename"],
    "networks": ["network", "networks", "cidr", "subnet", "subnets", "ip range",
                 "iprange", "ip_range", "address", "addresses", "ip", "ips",
                 "subnet(s)", "cidr(s)", "network(s)"],
    "level": ["level", "trust level", "trustlevel", "risk level"],
    "security_profile": ["security profile", "securityprofile", "profile",
                          "security_profile"],
    "force_assigned_vlan_id": ["force assigned vlan id", "force vlan",
                                "force_assigned_vlan", "forceassignedvlanid"],
    "mac_matching_fallback": ["mac matching fallback", "mac fallback",
                               "macmatchingfallback"],
    "extended_network_statistic_enabled": ["extended network statistic enabled",
                                            "extended stats", "extended statistics",
                                            "network statistics"],
    "use_label_as_node_device_id": ["use label as node device id", "label as device id",
                                     "use label"],
    "matching_vlan_id": ["matching vlan id", "matching vlan", "matchingvlanid"],
    "assigned_vlan_id": ["assigned vlan id", "assigned vlan", "vlan id", "vlan",
                          "assignedvlanid"],
    "adaptive_learning": ["adaptive learning", "adaptivelearning"],
    "learning": ["learning"],
    "is_public": ["is public", "public", "ispublic"],
    "merging_zone": ["merging zone", "merge zone", "mergingzone"],
    "merging": ["merging", "merge"],
    "isolated": ["isolated", "isolation", "is isolated"],
    "output_mode": ["output mode", "mode", "row mode", "outputmode"],
}


def _normalize(s):
    return re.sub(r"[^a-z0-9]", "", str(s).lower())


def _best_column_scores(field_keys, synonyms, source_columns):
    """Score every (field, column) pair and return the assignment obtained
    by greedily taking the globally best-scoring pairs first (rather than
    resolving one field at a time in a fixed order). This matters when two
    fields could both fuzzy-match the same column - e.g. a lone "VLAN"
    column should go to the field it matches *exactly* (assigned_vlan_id)
    rather than being grabbed by a field that only matches it loosely
    (matching_vlan_id) just because that field happens to be considered
    first."""
    norm_source = {col: _normalize(col) for col in source_columns}
    candidates = []  # (score, field_order, key, col)
    for field_order, key in enumerate(field_keys):
        norm_candidates = [_normalize(c) for c in [key] + synonyms.get(key, [])]
        for col, ncol in norm_source.items():
            best_score = 0.0
            if ncol in norm_candidates:
                best_score = 1.0
            else:
                for nc in norm_candidates:
                    score = difflib.SequenceMatcher(None, ncol, nc).ratio()
                    if ncol and nc and (ncol in nc or nc in ncol):
                        score = max(score, 0.9)
                    best_score = max(best_score, score)
            if best_score >= 0.6:
                candidates.append((best_score, field_order, key, col))

    # Best matches first; ties broken by the field's declared order so
    # results stay deterministic.
    candidates.sort(key=lambda c: (-c[0], c[1]))

    mapping = {key: "" for key in field_keys}
    used_cols = set()
    assigned_fields = set()
    for score, _order, key, col in candidates:
        if key in assigned_fields or col in used_cols:
            continue
        mapping[key] = col
        used_cols.add(col)
        assigned_fields.add(key)
    return mapping


def auto_map_columns(source_columns):
    """Return dict: field_key -> best matching source column (or "" if none)."""
    return _best_column_scores(FIELD_KEYS + ["output_mode"], SYNONYMS, source_columns)


# --------------------------------------------------------------------------
# File loaders
# --------------------------------------------------------------------------

def load_table_file(path):
    """Load a csv/xls/xlsx/docx file into a pandas DataFrame of strings."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        df = pd.read_csv(path, dtype=str, keep_default_na=False)
    elif ext in (".xlsx", ".xlsm"):
        df = pd.read_excel(path, dtype=str, engine="openpyxl")
        df = df.fillna("")
    elif ext == ".xls":
        df = pd.read_excel(path, dtype=str)
        df = df.fillna("")
    elif ext == ".docx":
        if docx is None:
            raise RuntimeError(
                "python-docx is not installed; cannot read .docx files. "
                "Install it with: pip install python-docx"
            )
        document = docx.Document(path)
        if not document.tables:
            raise RuntimeError("No table was found in this Word document.")
        table = document.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        header, *data = rows
        df = pd.DataFrame(data, columns=header)
    elif ext in (".txt", ".tsv"):
        # try to sniff delimiter
        with open(path, "r", encoding="utf-8-sig", errors="replace") as fh:
            sample = fh.read(4096)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
            sep = dialect.delimiter
        except csv.Error:
            sep = "\t" if "\t" in sample else ","
        df = pd.read_csv(path, dtype=str, sep=sep, keep_default_na=False)
    else:
        raise RuntimeError(f"Unsupported file type: {ext}")
    df.columns = [str(c).strip() for c in df.columns]
    df = df.astype(str)
    df = df.replace("nan", "")
    return df


def apply_mapping(df, mapping):
    """Build the working rows (list[dict]) using the field mapping."""
    rows = []
    for _, src_row in df.iterrows():
        row = {}
        for key in FIELD_KEYS:
            col = mapping.get(key, "")
            val = str(src_row[col]).strip() if col and col in df.columns else ""
            if val.lower() == "nan":
                val = ""
            row[key] = val
        col = mapping.get("output_mode", "")
        row["output_mode"] = (str(src_row[col]).strip().lower()
                               if col and col in df.columns else "")
        rows.append(row)
    return rows


# --------------------------------------------------------------------------
# Guardian config generation / parsing
# --------------------------------------------------------------------------

def generate_config_text(rows, default_mode="full"):
    lines = []
    for row in rows:
        zone = (row.get("zone_name") or "").strip()
        networks = (row.get("networks") or "").strip()
        if not zone:
            continue
        mode = (row.get("output_mode") or "").strip().lower() or default_mode
        if mode not in ("full", "compact"):
            mode = default_mode

        lines.append(f"vi zones add {networks} {zone}")

        for key, label, cmd, always_full, never_if_empty in FIELDS:
            if cmd is None:
                continue
            val = (row.get(key) or "").strip()
            if never_if_empty and not val:
                continue
            if mode == "compact" and not val:
                continue
            lines.append(f"vi zones {cmd} {val} {zone}")
    return "\n".join(lines) + ("\n" if lines else "")


def parse_config_text(text):
    """Parse an existing Guardian 'vi zones ...' config back into rows.
    Useful for re-opening / bulk-editing an already exported file."""
    cmd_to_key = {v: k for k, v in FIELD_CMDS.items()}
    rows_by_zone = {}
    order = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line.startswith("vi zones "):
            continue
        rest = line[len("vi zones "):]
        parts = rest.split(" ")
        if not parts:
            continue
        action = parts[0]
        if action == "add":
            body = rest[len("add "):]
            tokens = body.rsplit(" ", 1)
            if len(tokens) != 2:
                continue
            networks, zone = tokens[0].strip(), tokens[1].strip()
            if zone not in rows_by_zone:
                rows_by_zone[zone] = {k: "" for k in FIELD_KEYS}
                rows_by_zone[zone]["output_mode"] = ""
                order.append(zone)
            rows_by_zone[zone]["zone_name"] = zone
            rows_by_zone[zone]["networks"] = networks
        elif action in cmd_to_key:
            key = cmd_to_key[action]
            body = rest[len(action) + 1:]
            tokens = body.rsplit(" ", 1)
            if len(tokens) != 2:
                continue
            value, zone = tokens[0].strip(), tokens[1].strip()
            if zone not in rows_by_zone:
                rows_by_zone[zone] = {k: "" for k in FIELD_KEYS}
                rows_by_zone[zone]["output_mode"] = ""
                order.append(zone)
            rows_by_zone[zone][key] = value
    return [rows_by_zone[z] for z in order]


# --------------------------------------------------------------------------
# Suggest zones from a Guardian Asset Export
# --------------------------------------------------------------------------
# Guardian's Asset export (Assets page > Export) typically has one row per
# discovered device with columns for its IP address, VLAN ID and, often, a
# "Segment"/"Subnet"/"Network" label. This section groups that raw asset
# list into candidate zones: one zone per distinct segment (or, if no
# segment column exists, one zone per distinct VLAN with its IP addresses
# summarized into the smallest set of covering CIDR blocks).
#
# VLAN field rule (per Nozomi Guardian zone semantics):
#   - assigned_vlan_id: used in the common case, to *tag/assign* a VLAN to a
#     zone whose network range is unique.
#   - matching_vlan_id: used only to disambiguate zones whose IP range is
#     NOT unique - i.e. the same subnet/CIDR is reused across more than one
#     VLAN/segment (overlapping private ranges at different sites). In that
#     situation Guardian needs the VLAN tag *in addition to* the IP range to
#     match traffic to the correct zone, so matching_vlan_id is populated
#     and assigned_vlan_id is left blank.
ASSET_SYNONYMS = {
    "ip": ["ip", "ip address", "ipaddress", "address", "host", "host ip", "asset ip"],
    "vlan": ["vlan", "vlan id", "vlanid", "vlan tag"],
    "segment": ["segment", "subnet", "network", "zone", "site", "location",
                "segment name", "subnet name", "cidr"],
    "label": ["label", "name", "asset name", "hostname", "device name"],
}


def auto_map_asset_columns(source_columns):
    return _best_column_scores(list(ASSET_SYNONYMS.keys()), ASSET_SYNONYMS, source_columns)


def _sanitize_zone_name(name):
    name = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(name).strip())
    name = re.sub(r"_+", "_", name).strip("_")
    return name or "Zone"


def _looks_like_cidr(value):
    value = str(value).strip()
    if not value:
        return False
    try:
        import ipaddress
        ipaddress.ip_network(value, strict=False)
        return True
    except ValueError:
        return False


def _summarize_ips_to_cidr(ip_values):
    """Collapse a list of individual IPs/CIDRs into the smallest set of
    covering CIDR blocks."""
    import ipaddress
    nets = []
    for v in ip_values:
        v = str(v).strip()
        if not v:
            continue
        try:
            if "/" in v:
                nets.append(ipaddress.ip_network(v, strict=False))
            else:
                nets.append(ipaddress.ip_network(v + "/32", strict=False))
        except ValueError:
            continue
    if not nets:
        return ""
    collapsed = list(ipaddress.collapse_addresses(nets))
    return ",".join(str(n) for n in collapsed)


def suggest_zones_from_assets(df, mapping):
    """Group a raw asset export into candidate Guardian zones.

    Returns (rows, warnings) where rows matches the EDITABLE_FIELD_KEYS
    schema (ready to drop into the main editable grid / exporter) and
    warnings is a list of human-readable strings explaining assumptions
    made (e.g. which zones got matching_vlan_id because their network was
    not unique).
    """
    ip_col = mapping.get("ip", "")
    vlan_col = mapping.get("vlan", "")
    seg_col = mapping.get("segment", "")
    if not ip_col and not seg_col:
        raise ValueError("Map at least an IP Address column or a Segment/Subnet column.")

    # bucket assets: key = segment label if available, else vlan id
    buckets = {}   # bucket_key -> {"ips": [...], "vlans": set(), "segment_label": str}
    order = []
    for _, r in df.iterrows():
        ip_val = str(r[ip_col]).strip() if ip_col and ip_col in df.columns else ""
        vlan_val = str(r[vlan_col]).strip() if vlan_col and vlan_col in df.columns else ""
        seg_val = str(r[seg_col]).strip() if seg_col and seg_col in df.columns else ""

        if seg_val:
            key = seg_val
        elif vlan_val:
            key = f"VLAN{vlan_val}"
        else:
            continue

        if key not in buckets:
            buckets[key] = {"ips": [], "vlans": set(), "segment_label": seg_val}
            order.append(key)
        if ip_val:
            buckets[key]["ips"].append(ip_val)
        if vlan_val:
            buckets[key]["vlans"].add(vlan_val)

    # derive the network for each bucket
    bucket_networks = {}
    for key in order:
        b = buckets[key]
        if b["segment_label"] and _looks_like_cidr(b["segment_label"]):
            network = b["segment_label"]
        elif b["ips"]:
            network = _summarize_ips_to_cidr(b["ips"])
        else:
            network = ""
        bucket_networks[key] = network

    # find networks that are reused across more than one bucket/VLAN
    # (this is the "duplicate IP range" case that requires matching_vlan_id)
    network_to_keys = {}
    for key, net in bucket_networks.items():
        if net:
            network_to_keys.setdefault(net, []).append(key)
    duplicate_networks = {net for net, keys in network_to_keys.items() if len(keys) > 1}

    rows = []
    warnings = []
    used_names = set()
    for key in order:
        b = buckets[key]
        network = bucket_networks[key]
        vlans = sorted(b["vlans"])
        vlan_val = vlans[0] if len(vlans) == 1 else (",".join(vlans) if vlans else "")

        base_name = _sanitize_zone_name(b["segment_label"] or key)
        name = base_name
        n = 2
        while name in used_names:
            name = f"{base_name}_{n}"
            n += 1
        used_names.add(name)

        row = {k: "" for k in EDITABLE_FIELD_KEYS}
        row["zone_name"] = name
        row["networks"] = network
        row["output_mode"] = "compact"

        if network and network in duplicate_networks:
            # Same network appears under >1 VLAN/segment -> disambiguate via
            # matching_vlan_id, per the rule: matching_vlan_id is for
            # matching on segment+VLAN when duplicate IP ranges exist.
            row["matching_vlan_id"] = vlan_val
            warnings.append(
                f"Zone '{name}': network {network} is also used by another "
                f"suggested zone, so matching_vlan_id was set to '{vlan_val}' "
                f"instead of assigned_vlan_id to disambiguate them."
            )
        elif vlan_val:
            # Unique network -> the common case: assign the VLAN directly.
            row["assigned_vlan_id"] = vlan_val
            row["force_assigned_vlan_id"] = "true"

        if len(vlans) > 1:
            warnings.append(
                f"Zone '{name}' contains assets tagged with multiple VLANs "
                f"({', '.join(vlans)}); review whether it should be split "
                f"into separate zones."
            )
        if not network:
            warnings.append(f"Zone '{name}': could not derive a network/CIDR - please fill it in manually.")

        rows.append(row)

    return rows, warnings
